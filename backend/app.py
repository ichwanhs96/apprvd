# save this as app.py
from flask import Flask
import docx
import os

app = Flask(__name__)

@app.route("/")
def hello():
    return "Hello, World!"

@app.route("/documents")
def documents():
    current_dir = os.path.dirname(__file__)

    document = docx.Document(current_dir + '/' + 'Example docs.docx') # open an existing document

    print(document.paragraphs[5].text)
    print(document.tables[0].cell(0,0).text)

    p = document.tables[0].cell(0,0).add_paragraph('comment') # add a comment to a cell in a table
    
    comment = p.add_comment('comment',author='Obay Daba',initials= 'od') # add a comment on the entire paragraph

    paragraph1 = document.add_paragraph('text') # create new paragraph

    comment = paragraph1.add_comment('comment',author='Obay Daba',initials= 'od') # add a comment on the entire paragraph

    paragraph2 = document.add_paragraph('text') # create another paragraph

    run = paragraph2.add_run('texty') # add a run to the paragraph

    run.add_comment('comment') # add a comment only for the run text

    paragraph1.add_footnote('footnote text') # add a footnote

    document.save('new.docx') # save the document

    return "Document Updated"
